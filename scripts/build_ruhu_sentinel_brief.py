from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "competition"
DOCX_PATH = OUT_DIR / "ruhu-atlas-google-ai-agents-challenge-description.docx"
DIAGRAM_PATH = OUT_DIR / "ruhu-atlas-architecture.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1A1917"
MUTED = "5E6673"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
ORANGE = "E64E20"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text: str, bold=False, color=INK, size=11):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return run


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line=1.167)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix) :])
    else:
        add_run(p, text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4, line=1.167)
    add_run(p, text)


def add_callout(doc: Document, label: str, body: str, fill: str = CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0)
    add_run(p, f"{label}: ", bold=True, color=DARK_BLUE)
    add_run(p, body)
    doc.add_paragraph()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    widths = [2200, 7160]
    hdr = table.rows[0]
    hdr.cells[0].text = "Field"
    hdr.cells[1].text = "Recommendation"
    for cell in hdr.cells:
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            set_paragraph_spacing(p, after=0)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for key, value in rows:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                set_paragraph_spacing(p, after=0, line=1.10)
    doc.add_paragraph()


def pil_color(value: str) -> str:
    return value if value.startswith("#") else f"#{value}"


def draw_box(draw: ImageDraw.ImageDraw, xy, text, fill, outline, font, text_color=INK):
    fill = pil_color(fill)
    outline = pil_color(outline)
    text_color = pil_color(text_color)
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    words = text.split()
    lines = []
    line = ""
    for word in words:
        trial = (line + " " + word).strip()
        if draw.textlength(trial, font=font) <= (x2 - x1 - 36):
            line = trial
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    total_h = len(lines) * 24
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textlength(line, font=font)
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill=text_color)
        y += 24


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color=DARK_BLUE):
    color = pil_color(color)
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if ex > sx:
        points = [(ex, ey), (ex - 14, ey - 9), (ex - 14, ey + 9)]
    elif ex < sx:
        points = [(ex, ey), (ex + 14, ey - 9), (ex + 14, ey + 9)]
    elif ey > sy:
        points = [(ex, ey), (ex - 9, ey - 14), (ex + 9, ey - 14)]
    else:
        points = [(ex, ey), (ex - 9, ey + 14), (ex + 9, ey + 14)]
    draw.polygon(points, fill=color)


def make_diagram() -> None:
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 34)
        box_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 26)
        small_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 21)
    except OSError:
        title_font = ImageFont.load_default()
        box_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    draw.text((60, 40), "Ruhu Atlas: Build, Validate, Fix, and Operate Agents", fill=pil_color(DARK_BLUE), font=title_font)
    boxes = {
        "adk": (70, 160, 410, 300, "Gemini ADK Atlas Orchestrator", "#FFF3ED", ORANGE),
        "mcp": (540, 160, 880, 300, "MCP Tool Server", "#F4F6F9", DARK_BLUE),
        "api": (1010, 160, 1350, 300, "Ruhu API", "#E8EEF5", DARK_BLUE),
        "persona": (70, 430, 410, 570, "Workflow + Persona Generator", "#F9FAFB", MUTED),
        "runtime": (540, 430, 880, 570, "AgentDocument Runtime + Simulator", "#E8EEF5", DARK_BLUE),
        "trace": (1010, 430, 1350, 570, "Trace Store + Evaluation Rubrics", "#F9FAFB", MUTED),
        "patch": (300, 700, 640, 820, "Typed AgentDocument Deltas", "#FFF3ED", ORANGE),
        "report": (820, 700, 1160, 820, "Publish Readiness + Ops Report", "#E8EEF5", DARK_BLUE),
    }
    for x1, y1, x2, y2, text, fill, outline in boxes.values():
        draw_box(draw, (x1, y1, x2, y2), text, fill, outline, box_font)

    draw_arrow(draw, (410, 230), (540, 230), ORANGE)
    draw_arrow(draw, (880, 230), (1010, 230), DARK_BLUE)
    draw_arrow(draw, (240, 300), (240, 430), MUTED)
    draw_arrow(draw, (410, 500), (540, 500), DARK_BLUE)
    draw_arrow(draw, (880, 500), (1010, 500), DARK_BLUE)
    draw_arrow(draw, (1180, 430), (1180, 300), DARK_BLUE)
    draw_arrow(draw, (710, 570), (470, 700), ORANGE)
    draw_arrow(draw, (640, 760), (820, 760), DARK_BLUE)
    draw_arrow(draw, (1010, 760), (1180, 570), DARK_BLUE)

    draw.text(
        (60, 850),
        "Loop: understand workflow -> assemble AgentDocument -> simulate -> inspect traces -> fix -> rerun -> publish/operate.",
        fill=pil_color(MUTED),
        font=small_font,
    )
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    img.save(DIAGRAM_PATH)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(footer, "Ruhu Atlas project description | Google for Startups AI Agents Challenge", color=MUTED, size=9)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Ruhu Atlas", bold=True, color=DARK_BLUE, size=28)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    add_run(
        p,
        "Production Readiness Agent for African Customer-Support AI",
        color=MUTED,
        size=14,
    )

    add_callout(
        doc,
        "Recommended track",
        "Track 2: Optimize Existing Agents, with a credible Track 1 story for net-new agent assembly. Atlas already exists conceptually inside Ruhu as the authoring and operations copilot; this project expands it into a Gemini/ADK-powered agent that builds, validates, fixes, and operates Ruhu agents.",
        fill="FFF3ED",
    )
    add_key_value_table(
        doc,
        [
            ("Project name", "Ruhu Atlas: Production Readiness Agent for African Customer-Support AI"),
            ("Core idea", "Atlas is the agent inside Ruhu that turns rough business workflows into production-ready AgentDocuments, then validates, fixes, and monitors them."),
            ("African context", "Mixed-language support, WhatsApp-style text, voice-like utterances, unreliable integrations, fragmented channels, and high-volume customer operations."),
            ("Demo vertical", "Nigerian microfinance repayment and failed-payment support."),
            ("Submission assets", "Code, video, architecture diagram, test build/login, generated AgentDocument, simulation suite, and publish-readiness report."),
        ],
    )


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_diagram()
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)

    doc.add_heading("Project Description", level=1)
    p = doc.add_paragraph()
    add_run(
        p,
        "Ruhu Atlas is an AI agent inside the Ruhu platform that helps teams build, validate, fix, and operate deterministic customer-support agents for African markets.",
        bold=True,
        color=DARK_BLUE,
    )
    doc.add_paragraph(
        "The competition demo is intentionally narrow: Atlas builds and optimizes a Nigerian microfinance repayment-support agent. It uses Gemini, ADK, and MCP tools to inspect the workflow, propose typed AgentDocument changes, run realistic simulations, read traces, score reliability, patch failing flows, and generate a publish/no-publish decision. The runtime remains Ruhu's deterministic AgentDocument kernel; Atlas is the authoring and operations copilot that safely moves users from intent to deployment."
    )
    add_callout(
        doc,
        "Positioning",
        "Atlas is not a second runtime and not a generic chatbot builder. It is a structured agent that uses Ruhu's own tools to produce safer, more reliable customer-support agents.",
    )

    doc.add_heading("Design Principle", level=1)
    doc.add_paragraph(
        "Atlas should operate as a controlled copilot over typed platform artifacts. It proposes changes, explains tradeoffs, requests approvals, and invokes tools through MCP. It does not silently mutate production state or become the source of runtime truth."
    )
    add_key_value_table(
        doc,
        [
            ("Runtime truth", "AgentDocument remains the canonical runtime artifact: scenarios, steps, facts, transitions, tools, guards, handoff, and completion."),
            ("Atlas role", "Atlas owns planning, assembly assistance, validation, diagnosis, and operations recommendations."),
            ("Safety model", "Atlas emits typed deltas, blockers, questions, and reports. Human review gates apply before publish or risky changes."),
            ("Google role", "Gemini/ADK powers the orchestration loop, reasoning, simulation design, trace diagnosis, and report generation."),
        ],
    )

    doc.add_heading("Why This Fits Ruhu", level=1)
    for item in [
        "It matches the existing Atlas direction: authoring, tooling, validation, release, and operations copilot.",
        "It reinforces Ruhu's core thesis: deterministic conversation control plus bounded model assistance.",
        "It demonstrates production readiness, not just a working prototype.",
        "It is immediately useful for pilots because every enterprise buyer will ask how Ruhu proves agents behave correctly.",
        "It gives the judges visible autonomous behavior: understand, plan, assemble, simulate, inspect, diagnose, patch, rerun, and report.",
        "It creates a reusable internal product capability for future sales demos and customer deployments.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Product Modes", level=1)
    add_key_value_table(
        doc,
        [
            ("Build Mode", "Turns a rough workflow brief, PDF, API notes, or operator explanation into an AgentDocument plan and reviewable deltas."),
            ("Validate Mode", "Generates personas and edge cases, runs simulations, scores traces, and blocks risky publish attempts."),
            ("Fix Mode", "Diagnoses failing steps, transitions, facts, tools, or handoffs, then proposes typed AgentDocument patches."),
            ("Operate Mode", "Reads production traces, explains behavior, detects drift, and recommends improvements or regression tests."),
        ],
    )

    doc.add_heading("Core Capabilities", level=1)
    for item in [
        "Workflow understanding: convert business-process notes into scenario, step, fact, and tool requirements.",
        "AgentDocument assembly: propose typed deltas instead of prose-only edits.",
        "African-context simulation: generate Pidgin, mixed-language, low-literacy, angry, ambiguous, and channel-specific test personas.",
        "Trace diagnosis: inspect state transitions, facts, tool calls, semantic events, handoffs, and generated replies.",
        "Reliability scoring: produce measurable readiness scores for containment, safety, traceability, language robustness, and operational resilience.",
        "Patch-and-rerun loop: apply approved fixes, rerun failing cases, and show before/after improvement.",
        "Publish governance: produce a clear publish/no-publish report with blockers, residual risk, and required human approvals.",
        "Operations monitoring: turn live failures into new regression tests and improvement proposals.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Demo Scenario", level=1)
    doc.add_paragraph(
        "Use one concrete workflow: Atlas builds and optimizes a Nigerian microfinance repayment-support agent handling failed repayment, payment-not-reflected, loan-balance, repayment-plan, and escalation conversations."
    )
    add_key_value_table(
        doc,
        [
            ("Customer example", "\"I paid through Opay but una still dey call me.\""),
            ("Customer example", "\"My repayment no reflect.\""),
            ("Customer example", "\"I need small time, salary never enter.\""),
            ("Customer example", "\"I no understand the charges.\""),
            ("Customer example", "\"Transfer me to person abeg.\""),
        ],
    )
    doc.add_paragraph(
        "Atlas should help produce a Ruhu support agent that verifies identity, checks payment status, explains the balance, creates a ticket when payment is missing, offers an eligible repayment plan, and hands off safely when needed."
    )

    doc.add_heading("Before/After Demo Metrics", level=1)
    doc.add_paragraph(
        "Use measured or deterministic demo metrics in the video. The exact numbers can come from the local simulation suite, but the story should be before/after."
    )
    add_key_value_table(
        doc,
        [
            ("Test pass rate", "Before Atlas: 58%. After Atlas: 91%."),
            ("Unsafe replies", "Before Atlas: 5. After Atlas: 0."),
            ("Failed handoffs", "Before Atlas: 4. After Atlas: 1."),
            ("Tool-failure handling", "Before Atlas: 40%. After Atlas: 90%."),
            ("Pidgin/mixed-language completion", "Before Atlas: 45%. After Atlas: 85%."),
        ],
    )

    doc.add_heading("End-to-End Agent Loop", level=1)
    for step in [
        "Understands the business workflow and asks typed blocking questions when requirements are incomplete.",
        "Proposes an AgentDocument plan: scenarios, steps, fact schema, transitions, tools, guards, handoff, and completion.",
        "Generates realistic African customer personas and edge cases across English, Pidgin, mixed language, low-literacy phrasing, anger, missing account IDs, duplicate payments, tool failure, and escalation requests.",
        "Runs simulated conversations through Ruhu using MCP-exposed platform tools.",
        "Reads traces, semantic events, facts, tool calls, state transitions, and generated replies.",
        "Scores the run against a rubric: completion, safety, handoff quality, hallucination risk, tool behavior, language robustness, and customer outcome.",
        "Suggests or applies an AgentDocument patch, then reruns the failing cases.",
        "Produces a publish/no-publish report with before/after metrics, remaining risks, and recommended next actions.",
    ]:
        add_numbered(doc, step)

    doc.add_heading("Google Stack Fit", level=1)
    doc.add_paragraph(
        "Gemini and ADK should act as the Atlas orchestration brain. Ruhu should expose its deterministic runtime through MCP tools so Atlas can inspect, act, and iterate safely."
    )
    add_key_value_table(
        doc,
        [
            ("Gemini", "Workflow understanding, reasoning, diagnosis, persona generation, patch proposal, and final readiness report."),
            ("ADK", "Atlas orchestration: plan, call tools, evaluate results, and decide the next safe action."),
            ("MCP", "Secure bridge into Ruhu tools such as AgentDocument retrieval, simulation, trace retrieval, document patching, and evaluation."),
            ("Google Cloud", "Hosted demo, logs, optional Vertex/Gemini configuration, and production-style deployment story."),
        ],
    )

    doc.add_heading("Proposed MCP Tools", level=1)
    for tool in [
        "get_agent_document(agent_id): returns the current AgentDocument.",
        "propose_agent_document_deltas(goal, context): returns typed, reviewable changes.",
        "run_simulation(agent_id, persona, transcript_seed): starts a simulated conversation.",
        "get_trace(conversation_id): returns state, facts, tool calls, transitions, and replies.",
        "score_trace(trace, rubric): scores a trace against production criteria.",
        "patch_agent_document(agent_id, patch): applies a bounded document patch.",
        "rerun_eval_suite(agent_id): reruns known failing and regression cases.",
        "create_publish_report(agent_id): produces the readiness report.",
    ]:
        add_bullet(doc, tool)

    doc.add_heading("Architecture Diagram", level=1)
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Scoring Rubric", level=1)
    add_key_value_table(
        doc,
        [
            ("Containment", "Does the agent complete the target workflow without unnecessary human handoff?"),
            ("Safety", "Does it avoid unsafe commitments, policy violations, and hallucinated actions?"),
            ("Traceability", "Can every outcome be explained through state, facts, tools, transitions, and generated response?"),
            ("Language robustness", "Does it handle English, Pidgin, mixed language, ambiguous phrasing, and emotional users?"),
            ("Operational readiness", "Can it survive tool failures, missing data, duplicate events, and escalation requests?"),
            ("Improvement loop", "Does Atlas show measurable before/after reliability improvement?"),
        ],
    )

    doc.add_heading("Competition Story", level=1)
    doc.add_paragraph(
        "The demo narrative should be simple: a team gives Atlas a rough African customer-support workflow. Atlas creates or improves the Ruhu agent, tests it against realistic local edge cases, finds failures from traces, proposes typed fixes, reruns the test suite, and produces a publish decision."
    )
    for item in [
        "Before: a rough workflow or existing agent mishandles mixed-language payment disputes and tool failures.",
        "During: Atlas generates personas, runs simulations, and identifies failing transitions or missing handoffs.",
        "After: Atlas proposes typed deltas, reruns tests, and improves publish-readiness score.",
        "Business result: faster enterprise deployment with lower risk, clearer compliance evidence, and reusable regression tests.",
    ]:
        add_bullet(doc, item, bold_prefix=item.split(":")[0] + ":")

    doc.add_heading("Build Plan for the Deadline", level=1)
    add_key_value_table(
        doc,
        [
            ("Day 1", "Finalize the Atlas demo workflow, rubric, and baseline or generated AgentDocument."),
            ("Day 2", "Implement MCP tools for get document, propose deltas, run simulation, get trace, and score trace."),
            ("Day 3", "Build the Gemini/ADK Atlas loop: understand, plan, simulate, inspect, diagnose."),
            ("Day 4", "Add patch proposal, bounded patch application, and human approval checkpoints."),
            ("Day 5", "Create before/after demo cases and publish-readiness report UI or generated report."),
            ("Day 6", "Host demo build, record video, and polish architecture diagram."),
            ("Day 7", "Submission QA: code repo, test access, video, write-up, and fallback demo script."),
        ],
    )

    doc.add_heading("Submission Headline", level=1)
    add_callout(
        doc,
        "Headline",
        "Ruhu Atlas: an AI agent inside Ruhu that builds, validates, fixes, and operates African customer-support agents before and after they go live.",
        fill="FFF3ED",
    )
    doc.add_paragraph(
        "The strongest version of this project is narrow, practical, and directly tied to Ruhu's company thesis: deterministic agents are only valuable if enterprises can build them quickly, prove they behave correctly, and keep improving them after deployment."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
